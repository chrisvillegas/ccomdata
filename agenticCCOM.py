import streamlit as st
import pandas as pd
import platform
import keyring
from keyring.backends import SecretService
from keyring.backends import macOS
from langchain.retrievers import MultiQueryRetriever
from langchain_core.prompts import MessagesPlaceholder
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_experimental.agents import create_pandas_dataframe_agent
from langchain_community.vectorstores import Chroma
from langchain.prompts import ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
import matplotlib.pyplot as plt
import logging
from io import BytesIO
import docx
from fpdf import FPDF

logging.basicConfig(level=logging.INFO)

# Initialize session state
if 'messages' not in st.session_state:
    st.session_state['messages'] = []  # Initialize messages as an empty list


# Function to detect platform and configure keyring accordingly
def configure_keyring():
    if platform.system() == "Darwin":
        keyring.set_keyring(macOS.Keyring())  # Correct macOS backend
    elif platform.system() == "Linux":
        keyring.set_keyring(SecretService.Keyring())  # Correct Ubuntu backend
    else:
        st.error("Unsupported platform for keyring authentication.")


# Function to retrieve OpenAI API key from macOS or Ubuntu's keyring
def get_openai_api_key():
    configure_keyring()
    try:
        # Debug: Print out the current keyring backend
        st.write(f"Current keyring backend: {keyring.get_keyring()}")

        api_key = keyring.get_password("openai", "api_key")
        if not api_key:
            st.warning("No API key found. Please set your OpenAI API key.")
        return api_key
    except Exception as e:
        st.error(f"Failed to retrieve API key: {e}")
        return None


#def get_openai_api_key():
#    configure_keyring()
#    try:
#        api_key = keyring.get_password("openai", "api_key")
#        if not api_key:
#            st.warning("No API key found. Please set your OpenAI API key.")
#        return api_key
#    except Exception as e:
#        st.error(f"Failed to retrieve API key: {e}")
#        return None


# RAG chain setup for LangChain v0.3 and above
def setup_rag_chain(api_key, dfs):
    llm = ChatOpenAI(api_key=api_key, temperature=0, model="gpt-4o")
    embeddings = OpenAIEmbeddings(api_key=api_key)

    # Text splitter for creating document chunks
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

    documents = []
    for df in dfs:
        doc_text = df.to_string(index=False)
        documents.append(doc_text)

    split_docs = []
    for doc in documents:
        split_docs.extend(text_splitter.split_text(doc))

    # Store in vectorstore
    vectorstore = Chroma.from_texts(split_docs, embedding=embeddings)

    retriever = MultiQueryRetriever.from_llm(
        retriever=vectorstore.as_retriever(search_kwargs={"k": 10}),
        llm=llm
    )

    # Define the prompt using LangChain v0.3+ components
    prompt_template = ChatPromptTemplate.from_messages([
        SystemMessagePromptTemplate.from_template(
            "You are a Q&A assistant specializing in interpreting business performance data such as Nielsen ratings, advertising costs, and audience metrics. "
            "Use retrieved data to provide insights and suggest actions. If the query cannot be answered with the available data, reply 'No answer found.' "
            "When providing suggestions or insights, focus on improving business performance or optimizing costs based on trends."
        ),
        HumanMessagePromptTemplate.from_template(
            "Question: {input}\nContext: {context}\n\n"
            "Please interpret the provided data and suggest any actions or improvements if applicable."
        )
    ])

    # Retrieval chain setup
    qa_chain = create_retrieval_chain(
        retriever=retriever,
        combine_docs_chain=create_stuff_documents_chain(llm=llm, prompt=prompt_template)
    )

    return qa_chain


# Agent to interact with the dataframe, handling queries and plotting
def create_agent(api_key, df):
    llm = ChatOpenAI(api_key=api_key, temperature=0, model="gpt-4o")
    return create_pandas_dataframe_agent(llm=llm, df=df, allow_dangerous_code=True)


# Answer simple questions directly from the dataframe using an agent
def answer_simple_question(prompt, dfs):
    if "client name" in prompt.lower():
        client_name = extract_client_name(dfs)
        if client_name:
            return f"The client name is {client_name}."
        return "Client name not found in the data."

    if "date range" in prompt.lower():
        for df in dfs:
            if 'DATE' in df.columns or 'Date Range' in df.columns:
                date_range = df[['DATE']].dropna().iloc[0].to_string()
                return f"The date range is {date_range}."

    # Add plotting recommendations based on queries
    if "plot" in prompt.lower():
        generate_plot_based_on_query(prompt, dfs)  # Call the function to generate plot

    return None


# Function to generate plot based on query recommendation
def generate_plot_based_on_query(prompt, dfs):
    for df in dfs:
        # Example: Plot cost per station
        if "cost per station" in prompt.lower():
            plot_cost_per_station(df)
        # Example: Plot time series based on date and cost
        elif "time series" in prompt.lower() and 'DATE' in df.columns and 'COST' in df.columns:
            plot_time_series(df, 'DATE', 'COST')
        else:
            st.warning(f"Cannot find a valid plot based on your query: {prompt}")


# Function to plot cost per station
def plot_cost_per_station(df):
    if 'STATION' in df.columns and 'COST' in df.columns:
        df_grouped = df.groupby('STATION').sum()
        stations = df_grouped.index
        costs = df_grouped['COST']

        fig, ax = plt.subplots()
        ax.bar(stations, costs)
        ax.set_xlabel('TV Stations')
        ax.set_ylabel('Cost')
        ax.set_title('Cost per TV Station')
        st.pyplot(fig)  # This will display the plot in Streamlit


# Function to plot time series (e.g., DATE vs COST)
def plot_time_series(df, date_column, cost_column):
    df_filtered = df[[date_column, cost_column]].dropna()
    df_filtered[date_column] = pd.to_datetime(df_filtered[date_column])

    fig, ax = plt.subplots()
    ax.plot(df_filtered[date_column], df_filtered[cost_column], marker='o')
    ax.set_xlabel('Date')
    ax.set_ylabel('Cost')
    ax.set_title(f'{cost_column} over Time')
    st.pyplot(fig)  # This will display the plot in Streamlit


# Function to extract the client name from the dataframe
def extract_client_name(dfs):
    for df in dfs:
        for col in df.columns:
            if 'Client' in col:
                client_name = df[col].iloc[0]
                return client_name
    return None


# Function to format response data for display
def format_response_data(result):
    if isinstance(result, dict):
        if 'answer' in result:
            return f"**Answer:** {result['answer']}"
        elif 'output' in result:
            return f"**Answer:** {result['output']}"
        else:
            return "Sorry, I couldn't process that information."
    elif isinstance(result, str):
        return result.strip()
    elif isinstance(result, list):
        formatted_data = "\n".join([f"- {item}" for item in result])
        return formatted_data
    else:
        return str(result)


# Export chat history as Word
def export_word(chat_history):
    doc = docx.Document()
    for message in chat_history:
        role = message['role'].capitalize()
        content = message['content']

        # Add role as a heading, and the content as a paragraph
        doc.add_heading(f"{role}:", level=1)
        doc.add_paragraph(content)

    # Save the document to a BytesIO stream
    doc_output = BytesIO()
    doc.save(doc_output)
    doc_output.seek(0)
    return doc_output


# Export chat history as PDF
def export_pdf(chat_history):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for message in chat_history:
        pdf.multi_cell(200, 10, f"{message['role'].capitalize()}: {message['content']}")
    pdf_output = BytesIO()
    pdf.output(pdf_output, "F")
    pdf_output.seek(0)
    return pdf_output


# Main function for Streamlit app
def main():
    st.set_page_config(layout="wide", page_title="CCOM Data")
    st.title("CCOM Data")

    openai_api_key = get_openai_api_key()
    if not openai_api_key:
        st.write("Please enter an OpenAI API Key.")
        return

    dfs = None

    with st.sidebar:
        uploaded_file = st.file_uploader("Upload your Excel file", type=["xlsx"])
        if uploaded_file:
            with st.spinner("Processing your file..."):
                dfs = load_excel_file(uploaded_file)
                if dfs:
                    st.success("File processed.")

    if dfs:
        qa_chain = setup_rag_chain(openai_api_key, dfs)
        if qa_chain:
            agent = create_agent(openai_api_key, dfs[0])  # Create agent using the first DataFrame
            prompt = st.chat_input("Type your question here...")
            if prompt and isinstance(prompt, str) and prompt.strip() != "":
                st.session_state.messages.append({"role": "user", "content": prompt})

            with st.chat_message("user"):
                st.markdown(prompt)

            with st.chat_message("assistant"):
                message_placeholder = st.empty()
                with st.spinner('Generating response...'):
                    try:
                        # First, try to answer simple questions with the agent
                        response = answer_simple_question(prompt, dfs)
                        if response:
                            message_placeholder.markdown(f"**Answer:** {response}")
                        else:
                            # Use RAG for complex queries
                            response = qa_chain.invoke({"input": prompt})
                            message_placeholder.markdown(format_response_data(response))
                    except Exception as e:
                        response = f"An error occurred: {str(e)}"
                        message_placeholder.markdown(response)

                    st.session_state.messages.append({"role": "assistant", "content": response})

            # Export chat history buttons
            if st.session_state['messages']:
                if st.button("Export chat history as PDF"):
                    pdf_output = export_pdf(st.session_state['messages'])
                    st.download_button(
                        label="Download PDF",
                        data=pdf_output,
                        file_name="chat_history.pdf",
                        mime="application/pdf"
                    )
                if st.button("Export chat history as Word"):
                    word_output = export_word(st.session_state['messages'])
                    st.download_button(
                        label="Download Word",
                        data=word_output,
                        file_name="chat_history.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

# Function to load the Excel file
@st.cache_data
def load_excel_file(uploaded_file):
    try:
        dfs = pd.read_excel(uploaded_file, sheet_name=None, engine="openpyxl", header=None)
        processed_dfs = []
        total_sheets = len(dfs)

        for sheet_name, df in dfs.items():
            if len(df) < 16:
                st.sidebar.warning(f"Skipping sheet {sheet_name} because it has fewer than 16 rows.")
                continue

            for i in range(20):
                if set(df.iloc[i]).intersection(
                        {'LN#', 'PROGRAM', 'DATE', 'COST', 'EST (000)', 'ACT (000)', 'IDX'}):
                    df.columns = df.iloc[i]
                    df = df.drop(index=list(range(i + 1)))
                    df = rename_duplicate_columns(df)
                    df = df.fillna('')
                    processed_dfs.append(df)
                    break

        if not processed_dfs:
            st.sidebar.error("No usable data found in the uploaded file.")
            return None

        return processed_dfs

    except Exception as e:
        st.sidebar.error(f"Error processing Excel file: {e}")
        return None

# Function to rename duplicate columns
def rename_duplicate_columns(df):
    cols = pd.Series(df.columns)
    cols = cols.fillna('Unnamed')
    cols = cols.apply(lambda x: x if 'Unnamed' not in x else None)
    for dup in cols[cols.duplicated()].unique():
        dup_indices = cols[cols == dup].index.tolist()
        cols.iloc[dup_indices] = [f"{dup}_{i}" for i in range(1, len(dup_indices) + 1)]
    df.columns = cols
    return df.dropna(axis=1, how='all')

if __name__ == "__main__":
    main()

